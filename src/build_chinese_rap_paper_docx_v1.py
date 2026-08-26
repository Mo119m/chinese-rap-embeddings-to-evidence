from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "paper" / "manuscript.md"
FIGURES = ROOT / "figures"
OUTPUT = ROOT / "paper"
DOCX_PATH = OUTPUT / "Chinese_Rap_Evidence_Grounded_Manuscript.docx"

PAGE_WIDTH_DXA = 12240
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - (2 * MARGIN_DXA)

INK = "000000"
MUTED = "000000"
LIGHT = "E8EEF5"
RULE = "CBD5E1"


def set_rfonts(element, name: str) -> None:
    """Pin all script fonts and remove Word theme overrides."""
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), name)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attribute}")
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]


def set_run_font(run, name: str = "Times New Roman", size: float = 12, *, bold=None, italic=None, color=INK):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    set_rfonts(run._element, name)


def set_style_font(style, name: str, size: float, *, bold=None, italic=None, color=INK) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)
    set_rfonts(style._element, name)


def journal_heading(text: str) -> str:
    """Use an en space after a numbered heading without changing source Markdown."""
    return re.sub(r"^(\d+(?:\.\d+)*\.)\s+", lambda match: match.group(1) + "\u2002", text)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_geometry(table, widths: list[int]):
    """Set table, grid, column, and cell geometry to one explicit DXA contract."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    # Align the first cell's visible text with body text after its 90-DXA cell margin.
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")

    grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
    if len(grid_cols) != len(widths):
        raise ValueError("Unexpected table grid size")
    for idx, width in enumerate(widths):
        grid_cols[idx].set(qn("w:w"), str(width))


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_image_alt_text(inline_shape, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description[:80])


def add_markdown_runs(paragraph, text: str, *, base_size=12, base_color=INK):
    text = re.sub(r"\\\((.*?)\\\)", lambda m: m.group(1), text)
    text = re.sub(r"_\{([^}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^}]+)\}", r"^\1", text)
    for source, target in (
        (r"\ell", "ℓ"), (r"\alpha", "α"), (r"\sigma", "σ"), (r"\mu", "μ"),
        (r"\neq", "≠"), (r"\leq", "≤"), (r"\le", "≤"), (r"\geq", "≥"),
        (r"\mid", "|"), (r"\in", "∈"), (r"\times", "×"),
    ):
        text = text.replace(source, target)
    text = text.replace(r"\,", " ").replace(r"\{", "{").replace(r"\}", "}")
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=max(8, base_size - 0.5), color=base_color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, italic=True, color=base_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, "Times New Roman", 12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    set_style_font(h1, "Times New Roman", 13, bold=True, italic=False)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.line_spacing = 1.15
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, "Times New Roman", 12, bold=False, italic=True)
    h2.paragraph_format.space_before = Pt(9)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    for style_name in ("Caption",):
        style = styles[style_name]
        set_style_font(style, "Times New Roman", 10, bold=False, italic=False)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "Alt Text" not in [s.name for s in styles]:
        alt = styles.add_style("Alt Text", 1)
    else:
        alt = styles["Alt Text"]
    set_style_font(alt, "Times New Roman", 9, bold=False, italic=True, color=MUTED)
    alt.paragraph_format.left_indent = Inches(0.25)
    alt.paragraph_format.right_indent = Inches(0.25)
    alt.paragraph_format.line_spacing = 1.15
    alt.paragraph_format.space_after = Pt(6)

    if "References" not in [s.name for s in styles]:
        ref = styles.add_style("References", 1)
    else:
        ref = styles["References"]
    set_style_font(ref, "Times New Roman", 11, bold=False, italic=False)
    ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    ref.paragraph_format.left_indent = Inches(0.25)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.space_after = Pt(0)


def add_title_page(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("RESEARCH ARTICLE")
    set_run_font(run, size=9, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(title)
    set_run_font(run, size=19, bold=True, color=INK)

    for value, bold in (
        ("[Author name(s)]", True),
        ("[Affiliation(s)]", False),
        ("[Corresponding author email]", False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(value)
        set_run_font(run, size=11, bold=bold, color=INK if bold else MUTED)

    p.add_run().add_break(WD_BREAK.PAGE)


def add_document_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    set_run_font(run, size=17, bold=True)


def choose_column_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    if n == 2:
        return [int(CONTENT_WIDTH_DXA * 0.64), CONTENT_WIDTH_DXA - int(CONTENT_WIDTH_DXA * 0.64)]
    if n == 3:
        first = int(CONTENT_WIDTH_DXA * 0.46)
        rest = (CONTENT_WIDTH_DXA - first) // 2
        return [first, rest, CONTENT_WIDTH_DXA - first - rest]
    if n == 6:
        first = int(CONTENT_WIDTH_DXA * 0.34)
        rest = (CONTENT_WIDTH_DXA - first) // 5
        return [first, rest, rest, rest, rest, CONTENT_WIDTH_DXA - first - (4 * rest)]
    each = CONTENT_WIDTH_DXA // n
    return [each] * (n - 1) + [CONTENT_WIDTH_DXA - each * (n - 1)]


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    widths = choose_column_widths(rows[0])
    set_table_geometry(table, widths)

    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_together = True
            if c_idx > 0 and re.fullmatch(r"[\d.,/%–-]+|At least five songs per label|Effective-text mass at least 20", value.strip()):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_markdown_runs(p, value, base_size=9.4)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True

    # Subtle borders.
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), RULE)
        borders.append(node)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    spacer.paragraph_format.line_spacing = 1.0


def add_body_paragraph(doc: Document, text: str, *, after_heading=False, in_references=False):
    if in_references:
        p = doc.add_paragraph(style="References")
        add_markdown_runs(p, text, base_size=11)
        return p

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.widow_control = True
    is_abstract_field = bool(
        re.match(
            r"^\*\*(Purpose|Design/methodology/approach|Findings|Originality|Contribution to the field of Digital Humanities):",
            text,
        )
    )
    if not after_heading and not is_abstract_field and not re.match(r"^(RQ\d+\.|Keywords:|Table \d+\s+[A-Z]|Fig\. \d+\s+[A-Z]|Alt text:|\[)", text):
        p.paragraph_format.first_line_indent = Inches(0.25)
    match = re.match(r"^(RQ\d+\.)\s*(.*)$", text)
    if match:
        run = p.add_run(match.group(1) + " ")
        set_run_font(run, bold=True)
        add_markdown_runs(p, match.group(2))
    else:
        add_markdown_runs(p, text)
    return p


def add_math_block(doc: Document, latex: str):
    # The fusion equation also contains both z terms, so match it first.
    if "s_{q\\ell}^{F}" in latex:
        text = "sᶠ_qℓ = ½zᴰ_qℓ + ½zᴸ_qℓ."
    elif "z_{q\\ell}^{D}" in latex and "z_{q\\ell}^{L}" in latex:
        text = "zᴰ_qℓ = (sᴰ_qℓ − μᴰ_q) / σᴰ_q,    zᴸ_qℓ = (sᴸ_qℓ − μᴸ_q) / σᴸ_q."
    elif "P(y_t=k" in latex and "begin{cases}" in latex:
        text = "P(yₜ = k | xₜ) = g(xₜ), if k = c;    [1 − g(xₜ)]hₖ(xₜ), if k ≠ c."
    else:
        text = latex.replace("\\qquad", "    ").replace("\\,", " ").replace("&", "")
        text = re.sub(r"\\(?:begin|end)\{[^}]+\}", "", text)
        text = text.replace("\\", " ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name="Cambria Math", size=11.5, italic=True)


def build_docx(source: Path, docx_path: Path, figures: Path, *, include_title_page: bool, submission_mode: bool = False):
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    raw = source.read_text(encoding="utf-8")
    lines = raw.splitlines()
    title = lines[0].removeprefix("# ").strip()

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    section.different_first_page_header_footer = True
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    doc.core_properties.title = title
    doc.core_properties.subject = "Chinese rap lyrical-repertoire analysis"
    doc.core_properties.keywords = "Chinese rap, digital humanities, BGE-M3, lyrics, robustness"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""

    if include_title_page:
        add_title_page(doc, title)
        start = next(i for i, line in enumerate(lines) if line.strip() == "## Structured Abstract")
    else:
        add_document_title(doc, title)
        start = 1
    i = start
    after_heading = True
    in_references = False
    in_figure_legends = False
    last_shape = None
    expect_figure_caption = False
    expect_table_caption = False
    pending_figure_caption: str | None = None
    figure_legends: list[tuple[str, str]] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("## "):
            heading = journal_heading(stripped[3:])
            p = doc.add_paragraph(heading, style="Heading 1")
            if heading in {"References", "Figure Legends and Alt Text"}:
                p.paragraph_format.page_break_before = True
            in_references = heading == "References"
            in_figure_legends = heading == "Figure Legends and Alt Text"
            after_heading = True
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_paragraph(journal_heading(stripped[4:]), style="Heading 2")
            after_heading = True
            i += 1
            continue

        if stripped.startswith("[[FIGURE:"):
            filename = stripped.removeprefix("[[FIGURE:").removesuffix("]]")
            image_path = figures / filename
            if not image_path.exists():
                raise FileNotFoundError(image_path)
            if submission_mode:
                last_shape = None
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_together = True
                last_shape = p.add_run().add_picture(str(image_path), width=Inches(6.25))
            expect_figure_caption = True
            after_heading = False
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|\s*:?-+", lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = []
            for idx, table_line in enumerate(table_lines):
                cells = [c.strip() for c in table_line.strip("|").split("|")]
                if idx == 1 and all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
                    continue
                parsed.append(cells)
            add_table(doc, parsed)
            expect_table_caption = True
            after_heading = False
            continue

        if (expect_figure_caption and stripped.startswith("Fig. ")) or (
            expect_table_caption and stripped.startswith("Table ")
        ) or (in_figure_legends and stripped.startswith("Fig. ")):
            if submission_mode and expect_figure_caption:
                pending_figure_caption = stripped
            else:
                p = doc.add_paragraph(style="Caption")
                add_markdown_runs(p, stripped, base_size=10)
            expect_figure_caption = False
            expect_table_caption = False
            after_heading = False
            i += 1
            continue

        if stripped.startswith("Alt text:"):
            if submission_mode and pending_figure_caption is not None:
                figure_legends.append((pending_figure_caption, stripped))
                pending_figure_caption = None
            else:
                p = doc.add_paragraph(style="Alt Text")
                add_markdown_runs(p, stripped, base_size=9, base_color=MUTED)
                if last_shape is not None:
                    set_image_alt_text(last_shape, stripped.removeprefix("Alt text:").strip())
                    last_shape = None
            after_heading = False
            i += 1
            continue

        if stripped in {"\\[", "\\]"}:
            if stripped == "\\[":
                math_lines = []
                i += 1
                while i < len(lines) and lines[i].strip() != "\\]":
                    math_lines.append(lines[i].strip())
                    i += 1
                add_math_block(doc, " ".join(math_lines))
                i += 1
                after_heading = False
            else:
                i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_after = Pt(0)
            add_markdown_runs(p, stripped[2:])
            after_heading = False
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.35)
            p.paragraph_format.line_spacing = 1.5
            add_markdown_runs(p, stripped[2:], base_size=11)
            after_heading = False
            i += 1
            continue

        add_body_paragraph(doc, stripped, after_heading=after_heading, in_references=in_references)
        after_heading = False
        i += 1

    if submission_mode and figure_legends:
        doc.add_page_break()
        doc.add_paragraph("Figure Legends and Alt Text", style="Heading 1")
        for caption, alt_text in figure_legends:
            p = doc.add_paragraph(style="Caption")
            add_markdown_runs(p, caption, base_size=10)
            p = doc.add_paragraph(style="Alt Text")
            add_markdown_runs(p, alt_text, base_size=9, base_color=MUTED)

    # Keep document free of page and line numbering as required for the target submission format.
    for section in doc.sections:
        sect_pr = section._sectPr
        ln_num = sect_pr.find(qn("w:lnNumType"))
        if ln_num is not None:
            sect_pr.remove(ln_num)

    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=DOCX_PATH)
    parser.add_argument("--figures", type=Path, default=FIGURES)
    parser.add_argument("--no-title-page", action="store_true")
    parser.add_argument("--submission-mode", action="store_true", help="Omit embedded figures and collect legends/alt text on a final page")
    args = parser.parse_args()
    build_docx(args.source, args.output, args.figures, include_title_page=not args.no_title_page, submission_mode=args.submission_mode)
